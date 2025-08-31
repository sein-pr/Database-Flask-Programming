from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime, date, time, timedelta
import os
import json
from docx import Document
from docx.shared import Inches
import io

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///hospital.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create upload directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'doc', 'docx', 'xls', 'xlsx'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

db = SQLAlchemy(app)

# Database Models (3NF Normalized)
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    first_name = db.Column(db.String(50), nullable=False)
    last_name = db.Column(db.String(50), nullable=False)
    role = db.Column(db.String(20), default='user')  # admin, user
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class AccessRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(50), nullable=False)
    last_name = db.Column(db.String(50), nullable=False)
    reason = db.Column(db.Text, nullable=False)
    is_temporary = db.Column(db.Boolean, default=False)
    status = db.Column(db.String(20), default='pending')  # pending, approved, rejected
    requested_at = db.Column(db.DateTime, default=datetime.utcnow)

class Specialization(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True, nullable=False)
    doctors = db.relationship('Doctor', backref='specialization_ref', lazy=True)

class Doctor(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(50), nullable=False)
    last_name = db.Column(db.String(50), nullable=False)
    phone = db.Column(db.String(20))
    email = db.Column(db.String(100))
    specialization_id = db.Column(db.Integer, db.ForeignKey('specialization.id'), nullable=False)
    license_number = db.Column(db.String(50), unique=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    appointments = db.relationship('Appointment', backref='doctor_ref', lazy=True)

class Patient(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(50), nullable=False)
    last_name = db.Column(db.String(50), nullable=False)
    date_of_birth = db.Column(db.Date, nullable=False)
    gender = db.Column(db.String(10), nullable=False)
    phone = db.Column(db.String(20))
    email = db.Column(db.String(100))
    address = db.Column(db.Text)
    emergency_contact = db.Column(db.String(100))
    emergency_phone = db.Column(db.String(20))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    appointments = db.relationship('Appointment', backref='patient_ref', lazy=True)
    medical_records = db.relationship('MedicalRecord', backref='patient_ref', lazy=True)

    @property
    def age(self):
        today = date.today()
        return today.year - self.date_of_birth.year - ((today.month, today.day) < (self.date_of_birth.month, self.date_of_birth.day))

    @property
    def admission_count(self):
        return len(self.appointments)

class Appointment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    patient_id = db.Column(db.Integer, db.ForeignKey('patient.id'), nullable=False)
    doctor_id = db.Column(db.Integer, db.ForeignKey('doctor.id'), nullable=False)
    appointment_date = db.Column(db.Date, nullable=False)
    appointment_time = db.Column(db.Time, nullable=False)
    diagnosis = db.Column(db.Text)
    notes = db.Column(db.Text)
    status = db.Column(db.String(20), default='scheduled')  # scheduled, completed, cancelled
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class MedicalRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    patient_id = db.Column(db.Integer, db.ForeignKey('patient.id'), nullable=False)
    diagnosis = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    file_path = db.Column(db.String(255))
    file_name = db.Column(db.String(255))
    record_date = db.Column(db.Date, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Initialize database
with app.app_context():
    # Use instance folder for database (Flask default behavior)
    instance_db_path = os.path.join(app.instance_path, 'hospital.db')
    print(f"[v0] Database file path: {instance_db_path}")
    print(f"[v0] Database file exists: {os.path.exists(instance_db_path)}")
    
    # Ensure instance directory exists
    os.makedirs(app.instance_path, exist_ok=True)
    
    try:
        # Only create tables, don't recreate if they exist
        db.create_all()
        print("[v0] Database tables checked/created successfully")
        
        # Check if database is empty (first run)
        user_count = User.query.count()
        is_first_run = user_count == 0
        
        if is_first_run:
            print("[v0] First run detected - creating default data")
            
            # Create default admin user
            admin = User(
                username='admin',
                password_hash=generate_password_hash('admin123'),
                first_name='System',
                last_name='Administrator',
                role='admin'
            )
            db.session.add(admin)
            print("[v0] Created default admin user")
            
            # Create default specializations
            specializations = ['General Medicine', 'Cardiology', 'Neurology', 'Orthopedics', 'Pediatrics', 'Dermatology']
            for spec_name in specializations:
                spec = Specialization(name=spec_name)
                db.session.add(spec)
                print(f"[v0] Created specialization: {spec_name}")
            
            db.session.commit()
            print("[v0] Default data created successfully")
        else:
            print("[v0] Existing database found - preserving data")
        
        # Display current data counts
        user_count = User.query.count()
        patient_count = Patient.query.count()
        doctor_count = Doctor.query.count()
        appointment_count = Appointment.query.count()
        print(f"[v0] Current data counts - Users: {user_count}, Patients: {patient_count}, Doctors: {doctor_count}, Appointments: {appointment_count}")
        
    except Exception as e:
        print(f"[v0] Database initialization error: {str(e)}")
        raise

# Authentication routes
@app.route('/')
def index():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        user = User.query.filter_by(username=username, is_active=True).first()
        if user and check_password_hash(user.password_hash, password):
            session['user_id'] = user.id
            session['username'] = user.username
            session['role'] = user.role
            session['first_name'] = user.first_name
            session['last_name'] = user.last_name
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html')

@app.route('/request_access', methods=['GET', 'POST'])
def request_access():
    if request.method == 'POST':
        access_request = AccessRequest(
            first_name=request.form['first_name'],
            last_name=request.form['last_name'],
            reason=request.form['reason'],
            is_temporary=bool(request.form.get('is_temporary'))
        )
        db.session.add(access_request)
        db.session.commit()
        flash('Access request submitted successfully', 'success')
        return redirect(url_for('login'))
    
    return render_template('request_access.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Get dashboard statistics
    total_patients = Patient.query.count()
    total_doctors = Doctor.query.count()
    total_appointments = Appointment.query.count()
    pending_requests = AccessRequest.query.filter_by(status='pending').count()
    
    return render_template('dashboard.html', 
                         total_patients=total_patients,
                         total_doctors=total_doctors,
                         total_appointments=total_appointments,
                         pending_requests=pending_requests)

# Admin routes for managing access requests
@app.route('/admin/access_requests')
def admin_access_requests():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    pending_requests = AccessRequest.query.filter_by(status='pending').all()
    approved_requests = AccessRequest.query.filter_by(status='approved').all()
    rejected_requests = AccessRequest.query.filter_by(status='rejected').all()
    
    return render_template('admin/access_requests.html', 
                         pending_requests=pending_requests,
                         approved_requests=approved_requests,
                         rejected_requests=rejected_requests)

@app.route('/admin/approve_request/<int:request_id>')
def approve_request(request_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    access_request = AccessRequest.query.get_or_404(request_id)
    
    # Check if user already exists
    username = access_request.last_name[0].lower() + access_request.first_name.lower()
    existing_user = User.query.filter_by(username=username).first()
    
    if existing_user:
        # User exists, just reset password
        existing_user.password_hash = generate_password_hash('password@2025')
        existing_user.is_active = True
        flash(f'Password reset for user {username}. New password: password@2025', 'success')
    else:
        # Create new user
        new_user = User(
            username=username,
            password_hash=generate_password_hash('password@2025'),
            first_name=access_request.first_name,
            last_name=access_request.last_name,
            role='user'
        )
        db.session.add(new_user)
        flash(f'New user created: {username} with password: password@2025', 'success')
    
    # Update request status
    access_request.status = 'approved'
    db.session.commit()
    
    return redirect(url_for('admin_access_requests'))

@app.route('/admin/reject_request/<int:request_id>')
def reject_request(request_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    access_request = AccessRequest.query.get_or_404(request_id)
    access_request.status = 'rejected'
    db.session.commit()
    
    flash('Access request rejected', 'success')
    return redirect(url_for('admin_access_requests'))

@app.route('/admin/reset_password/<int:user_id>')
def reset_password(user_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    user.password_hash = generate_password_hash('password@2025')
    db.session.commit()
    
    flash(f'Password reset for {user.username}. New password: password@2025', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/users')
def admin_users():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    users = User.query.all()
    return render_template('admin/users.html', users=users)

@app.route('/admin/toggle_user/<int:user_id>')
def toggle_user(user_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    user = User.query.get_or_404(user_id)
    user.is_active = not user.is_active
    db.session.commit()
    
    status = 'activated' if user.is_active else 'deactivated'
    flash(f'User {user.username} has been {status}', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/dashboard')
def admin_dashboard():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('Access denied. Admin privileges required.', 'error')
        return redirect(url_for('dashboard'))
    
    # Calculate comprehensive statistics
    total_users = User.query.count()
    active_users = User.query.filter_by(is_active=True).count()
    total_patients = Patient.query.count()
    total_doctors = Doctor.query.count()
    total_appointments = Appointment.query.count()
    pending_requests = AccessRequest.query.filter_by(status='pending').count()
    total_records = MedicalRecord.query.count()
    specializations_count = Specialization.query.count()
    
    # Recent statistics (last 30 days)
    from datetime import datetime, timedelta
    thirty_days_ago = datetime.now() - timedelta(days=30)
    recent_patients = Patient.query.filter(Patient.created_at >= thirty_days_ago).count()
    
    # Today's appointments
    today = datetime.now().date()
    today_appointments = Appointment.query.filter(
        db.func.date(Appointment.appointment_date) == today
    ).count()
    
    # Recent records (last 7 days)
    seven_days_ago = datetime.now() - timedelta(days=7)
    recent_records = MedicalRecord.query.filter(MedicalRecord.created_at >= seven_days_ago).count()
    
    # Recent activity (mock data for demonstration)
    recent_activities = [
        {
            'action': 'New Patient Registered',
            'description': 'John Doe was added to the system',
            'timestamp': datetime.now() - timedelta(hours=2),
            'icon': 'user-plus'
        },
        {
            'action': 'Appointment Scheduled',
            'description': 'Dr. Smith scheduled with Jane Wilson',
            'timestamp': datetime.now() - timedelta(hours=4),
            'icon': 'calendar-plus'
        },
        {
            'action': 'Medical Record Updated',
            'description': 'Patient record updated with new diagnosis',
            'timestamp': datetime.now() - timedelta(hours=6),
            'icon': 'file-medical'
        }
    ]
    
    return render_template('admin/dashboard.html',
                         total_users=total_users,
                         active_users=active_users,
                         total_patients=total_patients,
                         recent_patients=recent_patients,
                         total_doctors=total_doctors,
                         specializations_count=specializations_count,
                         total_appointments=total_appointments,
                         today_appointments=today_appointments,
                         pending_requests=pending_requests,
                         total_records=total_records,
                         recent_activities=recent_activities)

# Patient management routes
@app.route('/patients')
def patients():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    search = request.args.get('search', '')
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    query = Patient.query
    if search:
        query = query.filter(
            (Patient.first_name.contains(search)) |
            (Patient.last_name.contains(search)) |
            (Patient.phone.contains(search)) |
            (Patient.email.contains(search))
        )
    
    patients = query.paginate(
        page=page, per_page=per_page, error_out=False
    )
    
    return render_template('patients/list.html', patients=patients, search=search)

@app.route('/patients/add', methods=['GET', 'POST'])
def add_patient():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        try:
            # Parse date of birth
            dob_str = request.form['date_of_birth']
            dob = datetime.strptime(dob_str, '%Y-%m-%d').date()
            
            patient = Patient(
                first_name=request.form['first_name'],
                last_name=request.form['last_name'],
                date_of_birth=dob,
                gender=request.form['gender'],
                phone=request.form.get('phone', ''),
                email=request.form.get('email', ''),
                address=request.form.get('address', ''),
                emergency_contact=request.form.get('emergency_contact', ''),
                emergency_phone=request.form.get('emergency_phone', '')
            )
            
            db.session.add(patient)
            db.session.commit()
            
            flash('Patient registered successfully!', 'success')
            return redirect(url_for('patients'))
            
        except ValueError as e:
            flash('Invalid date format. Please use YYYY-MM-DD format.', 'error')
        except Exception as e:
            flash('Error registering patient. Please try again.', 'error')
            db.session.rollback()
    
    return render_template('patients/add.html')

@app.route('/patients/<int:patient_id>')
def patient_detail(patient_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    patient = Patient.query.get_or_404(patient_id)
    recent_appointments = Appointment.query.filter_by(patient_id=patient_id)\
                                         .order_by(Appointment.appointment_date.desc())\
                                         .limit(5).all()
    recent_records = MedicalRecord.query.filter_by(patient_id=patient_id)\
                                       .order_by(MedicalRecord.record_date.desc())\
                                       .limit(5).all()
    
    return render_template('patients/detail.html', 
                         patient=patient, 
                         recent_appointments=recent_appointments,
                         recent_records=recent_records)

@app.route('/patients/<int:patient_id>/edit', methods=['GET', 'POST'])
def edit_patient(patient_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    patient = Patient.query.get_or_404(patient_id)
    
    if request.method == 'POST':
        try:
            # Parse date of birth
            dob_str = request.form['date_of_birth']
            dob = datetime.strptime(dob_str, '%Y-%m-%d').date()
            
            patient.first_name = request.form['first_name']
            patient.last_name = request.form['last_name']
            patient.date_of_birth = dob
            patient.gender = request.form['gender']
            patient.phone = request.form.get('phone', '')
            patient.email = request.form.get('email', '')
            patient.address = request.form.get('address', '')
            patient.emergency_contact = request.form.get('emergency_contact', '')
            patient.emergency_phone = request.form.get('emergency_phone', '')
            
            db.session.commit()
            
            flash('Patient information updated successfully!', 'success')
            return redirect(url_for('patient_detail', patient_id=patient.id))
            
        except ValueError as e:
            flash('Invalid date format. Please use YYYY-MM-DD format.', 'error')
        except Exception as e:
            flash('Error updating patient information. Please try again.', 'error')
            db.session.rollback()
    
    return render_template('patients/edit.html', patient=patient)

@app.route('/patients/<int:patient_id>/delete', methods=['POST'])
def delete_patient(patient_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    patient = Patient.query.get_or_404(patient_id)
    
    try:
        # Delete related records first
        MedicalRecord.query.filter_by(patient_id=patient_id).delete()
        Appointment.query.filter_by(patient_id=patient_id).delete()
        
        # Delete patient
        db.session.delete(patient)
        db.session.commit()
        
        flash('Patient deleted successfully!', 'success')
    except Exception as e:
        flash('Error deleting patient. Please try again.', 'error')
        db.session.rollback()
    
    return redirect(url_for('patients'))

# Doctor management routes
@app.route('/doctors')
def doctors():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    search = request.args.get('search', '')
    specialization_filter = request.args.get('specialization', '')
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    query = Doctor.query.join(Specialization)
    if search:
        query = query.filter(
            (Doctor.first_name.contains(search)) |
            (Doctor.last_name.contains(search)) |
            (Doctor.phone.contains(search)) |
            (Doctor.email.contains(search)) |
            (Doctor.license_number.contains(search))
        )
    
    if specialization_filter:
        query = query.filter(Doctor.specialization_id == specialization_filter)
    
    doctors = query.paginate(
        page=page, per_page=per_page, error_out=False
    )
    
    specializations = Specialization.query.all()
    
    return render_template('doctors/list.html', 
                         doctors=doctors, 
                         search=search,
                         specialization_filter=specialization_filter,
                         specializations=specializations)

@app.route('/doctors/add', methods=['GET', 'POST'])
def add_doctor():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        try:
            doctor = Doctor(
                first_name=request.form['first_name'],
                last_name=request.form['last_name'],
                phone=request.form.get('phone', ''),
                email=request.form.get('email', ''),
                specialization_id=request.form['specialization_id'],
                license_number=request.form.get('license_number', '')
            )
            
            db.session.add(doctor)
            db.session.commit()
            
            flash('Doctor registered successfully!', 'success')
            return redirect(url_for('doctors'))
            
        except Exception as e:
            flash('Error registering doctor. Please check if license number is unique.', 'error')
            db.session.rollback()
    
    specializations = Specialization.query.all()
    return render_template('doctors/add.html', specializations=specializations)

@app.route('/doctors/<int:doctor_id>')
def doctor_detail(doctor_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doctor = Doctor.query.get_or_404(doctor_id)
    recent_appointments = Appointment.query.filter_by(doctor_id=doctor_id)\
                                         .order_by(Appointment.appointment_date.desc())\
                                         .limit(10).all()
    
    # Get appointment statistics
    total_appointments = Appointment.query.filter_by(doctor_id=doctor_id).count()
    completed_appointments = Appointment.query.filter_by(doctor_id=doctor_id, status='completed').count()
    scheduled_appointments = Appointment.query.filter_by(doctor_id=doctor_id, status='scheduled').count()
    
    return render_template('doctors/detail.html', 
                         doctor=doctor, 
                         recent_appointments=recent_appointments,
                         total_appointments=total_appointments,
                         completed_appointments=completed_appointments,
                         scheduled_appointments=scheduled_appointments)

@app.route('/doctors/<int:doctor_id>/edit', methods=['GET', 'POST'])
def edit_doctor(doctor_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doctor = Doctor.query.get_or_404(doctor_id)
    
    if request.method == 'POST':
        try:
            doctor.first_name = request.form['first_name']
            doctor.last_name = request.form['last_name']
            doctor.phone = request.form.get('phone', '')
            doctor.email = request.form.get('email', '')
            doctor.specialization_id = request.form['specialization_id']
            doctor.license_number = request.form.get('license_number', '')
            
            db.session.commit()
            
            flash('Doctor information updated successfully!', 'success')
            return redirect(url_for('doctor_detail', doctor_id=doctor.id))
            
        except Exception as e:
            flash('Error updating doctor information. Please check if license number is unique.', 'error')
            db.session.rollback()
    
    specializations = Specialization.query.all()
    return render_template('doctors/edit.html', doctor=doctor, specializations=specializations)

@app.route('/doctors/<int:doctor_id>/delete', methods=['POST'])
def delete_doctor(doctor_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doctor = Doctor.query.get_or_404(doctor_id)
    
    try:
        # Check if doctor has appointments
        appointment_count = Appointment.query.filter_by(doctor_id=doctor_id).count()
        if appointment_count > 0:
            flash('Cannot delete doctor with existing appointments. Please reassign or cancel appointments first.', 'error')
            return redirect(url_for('doctor_detail', doctor_id=doctor_id))
        
        # Delete doctor
        db.session.delete(doctor)
        db.session.commit()
        
        flash('Doctor deleted successfully!', 'success')
    except Exception as e:
        flash('Error deleting doctor. Please try again.', 'error')
        db.session.rollback()
    
    return redirect(url_for('doctors'))

# Appointment management routes
@app.route('/appointments')
def appointments():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    view = request.args.get('view', 'list')  # list or calendar
    date_filter = request.args.get('date', '')
    doctor_filter = request.args.get('doctor', '')
    patient_filter = request.args.get('patient', '')
    status_filter = request.args.get('status', '')
    
    query = Appointment.query.join(Patient).join(Doctor)
    
    if date_filter:
        try:
            filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
            query = query.filter(Appointment.appointment_date == filter_date)
        except ValueError:
            pass
    
    if doctor_filter:
        query = query.filter(Appointment.doctor_id == doctor_filter)
    
    if patient_filter:
        query = query.filter(Appointment.patient_id == patient_filter)
    
    if status_filter:
        query = query.filter(Appointment.status == status_filter)
    
    appointments = query.order_by(Appointment.appointment_date.desc(), 
                                Appointment.appointment_time.desc()).all()
    
    doctors = Doctor.query.all()
    patients = Patient.query.all()
    
    return render_template('appointments/list.html', 
                         appointments=appointments,
                         doctors=doctors,
                         patients=patients,
                         view=view,
                         date_filter=date_filter,
                         doctor_filter=doctor_filter,
                         patient_filter=patient_filter,
                         status_filter=status_filter)

@app.route('/appointments/add', methods=['GET', 'POST'])
def add_appointment():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        try:
            appointment_date = datetime.strptime(request.form['appointment_date'], '%Y-%m-%d').date()
            appointment_time = datetime.strptime(request.form['appointment_time'], '%H:%M').time()
            
            # Check for conflicts
            existing = Appointment.query.filter_by(
                doctor_id=request.form['doctor_id'],
                appointment_date=appointment_date,
                appointment_time=appointment_time
            ).first()
            
            if existing:
                flash('This time slot is already booked for the selected doctor.', 'error')
                doctors = Doctor.query.all()
                patients = Patient.query.all()
                min_date = (date.today() + timedelta(days=1)).strftime('%Y-%m-%d')
                return render_template('appointments/add.html', doctors=doctors, patients=patients, min_date=min_date)
            
            appointment = Appointment(
                patient_id=request.form['patient_id'],
                doctor_id=request.form['doctor_id'],
                appointment_date=appointment_date,
                appointment_time=appointment_time,
                diagnosis=request.form.get('diagnosis', ''),
                notes=request.form.get('notes', ''),
                status='scheduled'
            )
            
            db.session.add(appointment)
            db.session.commit()
            
            flash('Appointment scheduled successfully!', 'success')
            return redirect(url_for('appointments'))
            
        except ValueError as e:
            flash('Invalid date or time format.', 'error')
        except Exception as e:
            flash('Error scheduling appointment. Please try again.', 'error')
            db.session.rollback()
    
    doctors = Doctor.query.all()
    patients = Patient.query.all()
    min_date = (date.today() + timedelta(days=1)).strftime('%Y-%m-%d')
    return render_template('appointments/add.html', doctors=doctors, patients=patients, min_date=min_date)

@app.route('/appointments/<int:appointment_id>')
def appointment_detail(appointment_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    appointment = Appointment.query.get_or_404(appointment_id)
    return render_template('appointments/detail.html', appointment=appointment)

@app.route('/appointments/<int:appointment_id>/edit', methods=['GET', 'POST'])
def edit_appointment(appointment_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    appointment = Appointment.query.get_or_404(appointment_id)
    
    if request.method == 'POST':
        try:
            appointment_date = datetime.strptime(request.form['appointment_date'], '%Y-%m-%d').date()
            appointment_time = datetime.strptime(request.form['appointment_time'], '%H:%M').time()
            
            # Check for conflicts (excluding current appointment)
            existing = Appointment.query.filter(
                Appointment.id != appointment_id,
                Appointment.doctor_id == request.form['doctor_id'],
                Appointment.appointment_date == appointment_date,
                Appointment.appointment_time == appointment_time
            ).first()
            
            if existing:
                flash('This time slot is already booked for the selected doctor.', 'error')
                doctors = Doctor.query.all()
                patients = Patient.query.all()
                return render_template('appointments/edit.html', 
                                     appointment=appointment, 
                                     doctors=doctors, 
                                     patients=patients)
            
            appointment.patient_id = request.form['patient_id']
            appointment.doctor_id = request.form['doctor_id']
            appointment.appointment_date = appointment_date
            appointment.appointment_time = appointment_time
            appointment.diagnosis = request.form.get('diagnosis', '')
            appointment.notes = request.form.get('notes', '')
            appointment.status = request.form['status']
            
            db.session.commit()
            
            flash('Appointment updated successfully!', 'success')
            return redirect(url_for('appointment_detail', appointment_id=appointment.id))
            
        except ValueError as e:
            flash('Invalid date or time format.', 'error')
        except Exception as e:
            flash('Error updating appointment. Please try again.', 'error')
            db.session.rollback()
    
    doctors = Doctor.query.all()
    patients = Patient.query.all()
    return render_template('appointments/edit.html', 
                         appointment=appointment, 
                         doctors=doctors, 
                         patients=patients)

@app.route('/appointments/<int:appointment_id>/delete', methods=['POST'])
def delete_appointment(appointment_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    appointment = Appointment.query.get_or_404(appointment_id)
    
    try:
        db.session.delete(appointment)
        db.session.commit()
        flash('Appointment deleted successfully!', 'success')
    except Exception as e:
        flash('Error deleting appointment. Please try again.', 'error')
        db.session.rollback()
    
    return redirect(url_for('appointments'))

# API endpoint for checking doctor availability
@app.route('/api/doctor-availability/<int:doctor_id>')
def doctor_availability(doctor_id):
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({'error': 'Date parameter required'}), 400
    
    try:
        check_date = datetime.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Invalid date format'}), 400
    
    # Get all booked times for this doctor on this date
    booked_appointments = Appointment.query.filter_by(
        doctor_id=doctor_id,
        appointment_date=check_date
    ).filter(Appointment.status != 'cancelled').all()
    
    booked_times = [apt.appointment_time.strftime('%H:%M') for apt in booked_appointments]
    
    # Generate available time slots (9 AM to 5 PM, 30-minute intervals)
    available_times = []
    current_time = time(9, 0)  # 9:00 AM
    end_time = time(17, 0)     # 5:00 PM
    
    while current_time < end_time:
        time_str = current_time.strftime('%H:%M')
        if time_str not in booked_times:
            available_times.append({
                'time': time_str,
                'display': current_time.strftime('%I:%M %p')
            })
        
        # Add 30 minutes
        current_datetime = datetime.combine(date.today(), current_time)
        current_datetime += timedelta(minutes=30)
        current_time = current_datetime.time()
    
    return jsonify({
        'available_times': available_times,
        'booked_times': booked_times
    })

# Calendar view API
@app.route('/api/calendar-appointments')
def calendar_appointments():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    start_date = request.args.get('start')
    end_date = request.args.get('end')
    
    if not start_date or not end_date:
        return jsonify({'error': 'Start and end dates required'}), 400
    
    try:
        start = datetime.strptime(start_date, '%Y-%m-%d').date()
        end = datetime.strptime(end_date, '%Y-%m-%d').date()
    except ValueError:
        return jsonify({'error': 'Invalid date format'}), 400
    
    appointments = Appointment.query.filter(
        Appointment.appointment_date >= start,
        Appointment.appointment_date <= end
    ).all()
    
    events = []
    for apt in appointments:
        events.append({
            'id': apt.id,
            'title': f"{apt.patient_ref.first_name} {apt.patient_ref.last_name} - Dr. {apt.doctor_ref.first_name} {apt.doctor_ref.last_name}",
            'start': f"{apt.appointment_date}T{apt.appointment_time}",
            'backgroundColor': '#50a69e' if apt.status == 'scheduled' else '#27ae60' if apt.status == 'completed' else '#e74c3c',
            'borderColor': '#073649',
            'textColor': '#ffffff'
        })
    
    return jsonify(events)

# Medical Records management routes
@app.route('/patients/<int:patient_id>/records')
def patient_records(patient_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    patient = Patient.query.get_or_404(patient_id)
    
    # Get filter parameters
    view_filter = request.args.get('filter', 'all')  # all, recent, date, diagnosis
    date_filter = request.args.get('date', '')
    diagnosis_filter = request.args.get('diagnosis', '')
    
    query = MedicalRecord.query.filter_by(patient_id=patient_id)
    
    if view_filter == 'recent':
        # Get records from last 30 days
        thirty_days_ago = date.today() - timedelta(days=30)
        query = query.filter(MedicalRecord.record_date >= thirty_days_ago)
    elif view_filter == 'date' and date_filter:
        try:
            filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
            query = query.filter(MedicalRecord.record_date == filter_date)
        except ValueError:
            pass
    elif view_filter == 'diagnosis' and diagnosis_filter:
        query = query.filter(MedicalRecord.diagnosis.contains(diagnosis_filter))
    
    records = query.order_by(MedicalRecord.record_date.desc()).all()
    
    # Get unique diagnoses for filter dropdown
    all_diagnoses = db.session.query(MedicalRecord.diagnosis.distinct())\
                             .filter_by(patient_id=patient_id)\
                             .all()
    diagnoses = [d[0] for d in all_diagnoses if d[0]]
    
    return render_template('patients/records.html', 
                         patient=patient, 
                         records=records,
                         diagnoses=diagnoses,
                         view_filter=view_filter,
                         date_filter=date_filter,
                         diagnosis_filter=diagnosis_filter)

@app.route('/patients/<int:patient_id>/records/add', methods=['GET', 'POST'])
def add_medical_record(patient_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    patient = Patient.query.get_or_404(patient_id)
    
    if request.method == 'POST':
        try:
            record_date = datetime.strptime(request.form['record_date'], '%Y-%m-%d').date()
            
            # Handle file upload
            file_path = None
            file_name = None
            if 'file' in request.files:
                file = request.files['file']
                if file and file.filename != '' and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    # Create unique filename
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    file_extension = filename.rsplit('.', 1)[1].lower()
                    unique_filename = f"patient_{patient_id}_{timestamp}.{file_extension}"
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                    file.save(file_path)
                    file_name = filename
            
            record = MedicalRecord(
                patient_id=patient_id,
                diagnosis=request.form['diagnosis'],
                description=request.form.get('description', ''),
                record_date=record_date,
                file_path=file_path,
                file_name=file_name
            )
            
            db.session.add(record)
            db.session.commit()
            
            flash('Medical record added successfully!', 'success')
            return redirect(url_for('patient_records', patient_id=patient_id))
            
        except ValueError as e:
            flash('Invalid date format.', 'error')
        except Exception as e:
            flash('Error adding medical record. Please try again.', 'error')
            db.session.rollback()
    
    return render_template('patients/add_record.html', patient=patient, date=date)

@app.route('/patients/<int:patient_id>/records/<int:record_id>/edit', methods=['GET', 'POST'])
def edit_medical_record(patient_id, record_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    patient = Patient.query.get_or_404(patient_id)
    record = MedicalRecord.query.get_or_404(record_id)
    
    if record.patient_id != patient_id:
        flash('Record not found for this patient.', 'error')
        return redirect(url_for('patient_records', patient_id=patient_id))
    
    if request.method == 'POST':
        try:
            record_date = datetime.strptime(request.form['record_date'], '%Y-%m-%d').date()
            
            # Handle file upload
            if 'file' in request.files:
                file = request.files['file']
                if file and file.filename != '' and allowed_file(file.filename):
                    # Delete old file if exists
                    if record.file_path and os.path.exists(record.file_path):
                        os.remove(record.file_path)
                    
                    filename = secure_filename(file.filename)
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    file_extension = filename.rsplit('.', 1)[1].lower()
                    unique_filename = f"patient_{patient_id}_{timestamp}.{file_extension}"
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                    file.save(file_path)
                    record.file_path = file_path
                    record.file_name = filename
            
            record.diagnosis = request.form['diagnosis']
            record.description = request.form.get('description', '')
            record.record_date = record_date
            
            db.session.commit()
            
            flash('Medical record updated successfully!', 'success')
            return redirect(url_for('patient_records', patient_id=patient_id))
            
        except ValueError as e:
            flash('Invalid date format.', 'error')
        except Exception as e:
            flash('Error updating medical record. Please try again.', 'error')
            db.session.rollback()
    
    return render_template('patients/edit_record.html', patient=patient, record=record)

@app.route('/patients/<int:patient_id>/records/<int:record_id>/delete', methods=['POST'])
def delete_medical_record(patient_id, record_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    record = MedicalRecord.query.get_or_404(record_id)
    
    if record.patient_id != patient_id:
        flash('Record not found for this patient.', 'error')
        return redirect(url_for('patient_records', patient_id=patient_id))
    
    try:
        # Delete file if exists
        if record.file_path and os.path.exists(record.file_path):
            os.remove(record.file_path)
        
        db.session.delete(record)
        db.session.commit()
        
        flash('Medical record deleted successfully!', 'success')
    except Exception as e:
        flash('Error deleting medical record. Please try again.', 'error')
        db.session.rollback()
    
    return redirect(url_for('patient_records', patient_id=patient_id))

@app.route('/download/<int:record_id>')
def download_file(record_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    record = MedicalRecord.query.get_or_404(record_id)
    
    if not record.file_path or not os.path.exists(record.file_path):
        flash('File not found.', 'error')
        return redirect(url_for('patient_records', patient_id=record.patient_id))
    
    return send_file(record.file_path, as_attachment=True, download_name=record.file_name)

# Export functionality
@app.route('/patients/<int:patient_id>/export')
def export_patient_records(patient_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    patient = Patient.query.get_or_404(patient_id)
    export_format = request.args.get('format', 'txt')  # txt or docx
    diagnosis_filter = request.args.get('diagnosis', '')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    
    # Build query
    query = MedicalRecord.query.filter_by(patient_id=patient_id)
    
    if diagnosis_filter:
        query = query.filter(MedicalRecord.diagnosis.contains(diagnosis_filter))
    
    if date_from:
        try:
            from_date = datetime.strptime(date_from, '%Y-%m-%d').date()
            query = query.filter(MedicalRecord.record_date >= from_date)
        except ValueError:
            pass
    
    if date_to:
        try:
            to_date = datetime.strptime(date_to, '%Y-%m-%d').date()
            query = query.filter(MedicalRecord.record_date <= to_date)
        except ValueError:
            pass
    
    records = query.order_by(MedicalRecord.record_date.desc()).all()
    
    if export_format == 'docx':
        return export_to_word(patient, records)
    else:
        return export_to_text(patient, records)

def export_to_text(patient, records):
    """Export patient records to text file"""
    output = io.StringIO()
    
    # Header information
    output.write("SEIN HOSPITAL MANAGEMENT SYSTEM\n")
    output.write("=" * 50 + "\n\n")
    output.write("PATIENT MEDICAL RECORDS EXPORT\n\n")
    
    # Patient information
    output.write(f"Patient Name: {patient.first_name} {patient.last_name}\n")
    output.write(f"Date of Birth: {patient.date_of_birth.strftime('%B %d, %Y')}\n")
    output.write(f"Age: {patient.age} years\n")
    output.write(f"Gender: {patient.gender}\n")
    if patient.phone:
        output.write(f"Phone: {patient.phone}\n")
    if patient.email:
        output.write(f"Email: {patient.email}\n")
    
    output.write(f"\nExport Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
    output.write(f"Total Records: {len(records)}\n")
    output.write("\n" + "=" * 50 + "\n\n")
    
    # Medical records
    if records:
        for i, record in enumerate(records, 1):
            output.write(f"RECORD #{i}\n")
            output.write("-" * 20 + "\n")
            output.write(f"Date: {record.record_date.strftime('%B %d, %Y')}\n")
            output.write(f"Diagnosis: {record.diagnosis}\n")
            if record.description:
                output.write(f"Description: {record.description}\n")
            if record.file_name:
                output.write(f"Attached File: {record.file_name}\n")
            output.write(f"Record Created: {record.created_at.strftime('%B %d, %Y at %I:%M %p')}\n")
            output.write("\n")
    else:
        output.write("No medical records found for the specified criteria.\n")
    
    # Create response
    output.seek(0)
    filename = f"{patient.first_name}_{patient.last_name}_medical_records_{datetime.now().strftime('%Y%m%d')}.txt"
    
    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8')),
        mimetype='text/plain',
        as_attachment=True,
        download_name=filename
    )

def export_to_word(patient, records):
    """Export patient records to Word document"""
    doc = Document()
    
    # Header
    header = doc.add_heading('SEIN HOSPITAL MANAGEMENT SYSTEM', 0)
    header.alignment = 1  # Center alignment
    
    doc.add_heading('Patient Medical Records Export', level=1)
    
    # Patient information table
    patient_table = doc.add_table(rows=6, cols=2)
    patient_table.style = 'Table Grid'
    
    patient_info = [
        ('Patient Name', f"{patient.first_name} {patient.last_name}"),
        ('Date of Birth', patient.date_of_birth.strftime('%B %d, %Y')),
        ('Age', f"{patient.age} years"),
        ('Gender', patient.gender),
        ('Phone', patient.phone or 'Not provided'),
        ('Email', patient.email or 'Not provided')
    ]
    
    for i, (label, value) in enumerate(patient_info):
        patient_table.cell(i, 0).text = label
        patient_table.cell(i, 1).text = value
    
    doc.add_paragraph()
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Total Records: {len(records)}")
    
    # Medical records
    if records:
        doc.add_heading('Medical Records', level=2)
        
        for i, record in enumerate(records, 1):
            doc.add_heading(f'Record #{i}', level=3)
            
            record_table = doc.add_table(rows=4 + (1 if record.file_name else 0), cols=2)
            record_table.style = 'Table Grid'
            
            record_info = [
                ('Date', record.record_date.strftime('%B %d, %Y')),
                ('Diagnosis', record.diagnosis),
                ('Description', record.description or 'Not provided'),
                ('Record Created', record.created_at.strftime('%B %d, %Y at %I:%M %p'))
            ]
            
            if record.file_name:
                record_info.append(('Attached File', record.file_name))
            
            for j, (label, value) in enumerate(record_info):
                record_table.cell(j, 0).text = label
                record_table.cell(j, 1).text = value
            
            doc.add_paragraph()
    else:
        doc.add_paragraph("No medical records found for the specified criteria.")
    
    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    filename = f"{patient.first_name}_{patient.last_name}_medical_records_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return send_file(
        doc_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )

if __name__ == '__main__':
    app.run(debug=True)
